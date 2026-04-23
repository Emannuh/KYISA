"""
KYISA Teams — Bulk Player Upload (CSV/Excel)

Allows team managers to upload a CSV/Excel file with player data
for bulk creation/update.

Endpoints:
  POST /api/v1/teams/players/bulk-upload/
  GET  /api/v1/teams/players/bulk-template/
"""
import csv
import io
import logging
from datetime import datetime

from rest_framework import status, permissions
from rest_framework.views import APIView
from rest_framework.response import Response
from django.http import HttpResponse
from drf_spectacular.utils import extend_schema

from accounts.models import KenyaCounty
from accounts.permissions import IsTeamManager, IsCompetitionManager
from competitions.models import SportType
from teams.models import Player, Team, CountyRegistration, CountyDiscipline, CountyPlayer

logger = logging.getLogger(__name__)

REQUIRED_COLUMNS = ["first_name", "last_name", "date_of_birth", "position", "shirt_number"]
OPTIONAL_COLUMNS = ["national_id_number", "birth_cert_number", "phone"]
ALL_COLUMNS = REQUIRED_COLUMNS + OPTIONAL_COLUMNS

VALID_POSITIONS = ["GK", "CB", "LB", "RB", "CDM", "CM", "AM", "LW", "RW", "CF", "ST"]


def _parse_csv(file_obj):
    """Parse a CSV file and return list of row dicts + errors."""
    try:
        content = file_obj.read().decode("utf-8-sig")
    except UnicodeDecodeError:
        content = file_obj.read().decode("latin-1")

    reader = csv.DictReader(io.StringIO(content))
    rows = []
    errors = []

    # Check required headers
    headers = [h.strip().lower().replace(" ", "_") for h in (reader.fieldnames or [])]
    missing = [c for c in REQUIRED_COLUMNS if c not in headers]
    if missing:
        return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

    for i, raw_row in enumerate(reader, start=2):  # row 2 = first data row
        row = {k.strip().lower().replace(" ", "_"): (v.strip() if v else "") for k, v in raw_row.items()}
        row_errors = []

        # Required field checks
        for col in REQUIRED_COLUMNS:
            if not row.get(col):
                row_errors.append(f"Missing '{col}'")

        # Validate date_of_birth
        dob = row.get("date_of_birth", "")
        parsed_dob = None
        if dob:
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
                try:
                    parsed_dob = datetime.strptime(dob, fmt).date()
                    break
                except ValueError:
                    continue
            if not parsed_dob:
                row_errors.append(f"Invalid date_of_birth '{dob}' (use YYYY-MM-DD or DD/MM/YYYY)")
        row["_parsed_dob"] = parsed_dob

        # Validate position
        pos = row.get("position", "").upper()
        if pos and pos not in VALID_POSITIONS:
            row_errors.append(f"Invalid position '{pos}'. Must be one of: {', '.join(VALID_POSITIONS)}")
        row["position"] = pos

        # Validate shirt_number
        shirt = row.get("shirt_number", "")
        if shirt:
            try:
                row["shirt_number"] = int(shirt)
            except ValueError:
                row_errors.append(f"Invalid shirt_number '{shirt}' (must be a number)")

        if row_errors:
            errors.append({"row": i, "errors": row_errors})
        rows.append(row)

    return rows, errors


def _parse_excel(file_obj):
    """Parse an Excel file and return list of row dicts + errors."""
    try:
        import openpyxl
    except ImportError:
        return [], [{"row": 0, "error": "openpyxl is required for Excel uploads. Install with: pip install openpyxl"}]

    try:
        wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        return [], [{"row": 0, "error": f"Cannot read Excel file: {e}"}]

    header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        return [], [{"row": 0, "error": "Empty spreadsheet"}]

    headers = [str(h or "").strip().lower().replace(" ", "_") for h in header_row]
    missing = [c for c in REQUIRED_COLUMNS if c not in headers]
    if missing:
        return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

    rows = []
    errors = []
    for i, excel_row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        raw = {headers[j]: (str(v).strip() if v is not None else "") for j, v in enumerate(excel_row) if j < len(headers)}
        row_errors = []

        for col in REQUIRED_COLUMNS:
            if not raw.get(col):
                row_errors.append(f"Missing '{col}'")

        # Parse DOB
        dob_val = raw.get("date_of_birth", "")
        parsed_dob = None
        if isinstance(excel_row[headers.index("date_of_birth")] if "date_of_birth" in headers else None, datetime):
            parsed_dob = excel_row[headers.index("date_of_birth")].date()
        elif dob_val:
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S"):
                try:
                    parsed_dob = datetime.strptime(dob_val, fmt).date()
                    break
                except ValueError:
                    continue
            if not parsed_dob:
                row_errors.append(f"Invalid date_of_birth '{dob_val}'")
        raw["_parsed_dob"] = parsed_dob

        pos = raw.get("position", "").upper()
        if pos and pos not in VALID_POSITIONS:
            row_errors.append(f"Invalid position '{pos}'")
        raw["position"] = pos

        shirt = raw.get("shirt_number", "")
        if shirt:
            try:
                raw["shirt_number"] = int(float(shirt))
            except (ValueError, TypeError):
                row_errors.append(f"Invalid shirt_number '{shirt}'")

        if row_errors:
            errors.append({"row": i, "errors": row_errors})
        rows.append(raw)

    wb.close()
    return rows, errors


class BulkPlayerUploadView(APIView):
    """
    POST /api/v1/teams/players/bulk-upload/
    Upload CSV or Excel file to bulk create players for a team.
    """
    permission_classes = [IsTeamManager]

    @extend_schema(tags=["teams"], summary="Bulk upload players via CSV/Excel")
    def post(self, request):
        team_id = request.data.get("team_id")
        file = request.FILES.get("file")

        if not file:
            return Response({"detail": "No file uploaded."}, status=status.HTTP_400_BAD_REQUEST)
        if not team_id:
            return Response({"detail": "team_id is required."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            team = Team.objects.get(pk=team_id)
        except Team.DoesNotExist:
            return Response({"detail": "Team not found."}, status=status.HTTP_404_NOT_FOUND)

        # Check ownership
        if team.manager != request.user:
            return Response({"detail": "You are not the manager of this team."}, status=status.HTTP_403_FORBIDDEN)

        # Determine file type and parse
        filename = file.name.lower()
        if filename.endswith(".csv"):
            rows, parse_errors = _parse_csv(file)
        elif filename.endswith((".xlsx", ".xls")):
            rows, parse_errors = _parse_excel(file)
        else:
            return Response(
                {"detail": "Unsupported file format. Use .csv or .xlsx"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if parse_errors:
            return Response({
                "detail": "File has validation errors. No players were created.",
                "errors": parse_errors,
                "total_rows": len(rows),
            }, status=status.HTTP_400_BAD_REQUEST)

        # Preview mode: dry_run=true returns parsed data without saving
        if request.data.get("dry_run") in ("true", "1", True):
            preview = []
            for row in rows:
                preview.append({
                    "first_name": row.get("first_name", ""),
                    "last_name": row.get("last_name", ""),
                    "date_of_birth": row.get("_parsed_dob", ""),
                    "position": row.get("position", ""),
                    "shirt_number": row.get("shirt_number", ""),
                    "national_id_number": row.get("national_id_number", ""),
                })
            return Response({
                "detail": f"Preview: {len(preview)} players parsed successfully.",
                "players": preview,
                "dry_run": True,
            })

        # Create players
        created = 0
        updated = 0
        row_errors = []
        for i, row in enumerate(rows, start=2):
            try:
                id_number = row.get("national_id_number", "").strip()
                defaults = {
                    "first_name": row["first_name"],
                    "last_name": row["last_name"],
                    "date_of_birth": row["_parsed_dob"],
                    "position": row["position"],
                    "shirt_number": row.get("shirt_number", 0),
                }
                if row.get("birth_cert_number"):
                    defaults["birth_cert_number"] = row["birth_cert_number"]

                if id_number:
                    player, was_created = Player.objects.update_or_create(
                        team=team,
                        national_id_number=id_number,
                        defaults=defaults,
                    )
                else:
                    player = Player.objects.create(team=team, **defaults)
                    was_created = True

                if was_created:
                    created += 1
                else:
                    updated += 1

            except Exception as exc:
                row_errors.append({"row": i, "error": str(exc)})

        result = {
            "detail": f"Bulk upload complete: {created} created, {updated} updated.",
            "created": created,
            "updated": updated,
            "total_rows": len(rows),
        }
        if row_errors:
            result["errors"] = row_errors

        return Response(result, status=status.HTTP_201_CREATED)


class BulkPlayerTemplateView(APIView):
    """
    GET /api/v1/teams/players/bulk-template/
    Download a CSV template for bulk player upload.
    """
    permission_classes = [permissions.IsAuthenticated]

    @extend_schema(tags=["teams"], summary="Download CSV template for bulk player upload")
    def get(self, request):
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = 'attachment; filename="kyisa_player_upload_template.csv"'

        writer = csv.writer(response)
        writer.writerow([
            "first_name", "last_name", "date_of_birth", "position",
            "shirt_number", "national_id_number", "birth_cert_number",
        ])
        # Example row
        writer.writerow([
            "John", "Doe", "2008-03-15", "CF", "9", "12345678", "BC-001234",
        ])
        writer.writerow([
            "Jane", "Smith", "2007-11-20", "GK", "1", "87654321", "BC-005678",
        ])

        return response


SECRETARY_REQUIRED_COLUMNS = ["name", "id_number", "county", "date_of_birth"]
SECRETARY_OPTIONAL_COLUMNS = ["discipline"]


def _norm_header(value):
    return str(value or "").strip().lower().replace(" ", "_")


def _extract_secretary_value(row, *aliases):
    for alias in aliases:
        if alias in row and str(row.get(alias, "")).strip():
            return str(row.get(alias, "")).strip()
    return ""


def _parse_date(date_text):
    if not date_text:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(date_text, fmt).date()
        except ValueError:
            continue
    return None


def _split_name(full_name):
    parts = [p for p in str(full_name or "").strip().split() if p]
    if not parts:
        return "", ""
    if len(parts) == 1:
        return parts[0], "-"
    return parts[0], " ".join(parts[1:])


def _parse_secretary_csv(file_obj):
    try:
        content = file_obj.read().decode("utf-8-sig")
    except UnicodeDecodeError:
        content = file_obj.read().decode("latin-1")

    reader = csv.DictReader(io.StringIO(content))
    headers = [_norm_header(h) for h in (reader.fieldnames or [])]
    missing = [c for c in SECRETARY_REQUIRED_COLUMNS if c not in headers]
    if missing:
        return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

    rows = []
    errors = []
    for i, raw_row in enumerate(reader, start=2):
        row = {_norm_header(k): (str(v).strip() if v is not None else "") for k, v in raw_row.items()}
        rows.append(row)
        if not row.get("name") or not row.get("id_number") or not row.get("county") or not row.get("date_of_birth"):
            errors.append({"row": i, "error": "Missing one or more required values: name, id_number, county, date_of_birth"})
    return rows, errors


def _parse_secretary_excel(file_obj):
    try:
        import openpyxl
    except ImportError:
        return [], [{"row": 0, "error": "openpyxl is required for Excel uploads."}]

    try:
        wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        return [], [{"row": 0, "error": f"Cannot read Excel file: {e}"}]

    header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        return [], [{"row": 0, "error": "Empty spreadsheet"}]

    headers = [_norm_header(h) for h in header_row]
    missing = [c for c in SECRETARY_REQUIRED_COLUMNS if c not in headers]
    if missing:
        wb.close()
        return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

    rows = []
    errors = []
    for i, excel_row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        row = {
            headers[j]: (str(v).strip() if v is not None else "")
            for j, v in enumerate(excel_row)
            if j < len(headers)
        }
        rows.append(row)
        if not row.get("name") or not row.get("id_number") or not row.get("county") or not row.get("date_of_birth"):
            errors.append({"row": i, "error": "Missing one or more required values: name, id_number, county, date_of_birth"})

    wb.close()
    return rows, errors


def _parse_secretary_word(file_obj):
    try:
        from docx import Document
    except ImportError:
        return [], [{"row": 0, "error": "python-docx is required for Word uploads (.docx)."}]

    try:
        document = Document(file_obj)
    except Exception as e:
        return [], [{"row": 0, "error": f"Cannot read Word document: {e}"}]

    rows = []
    errors = []

    # Preferred format: first table with headers.
    if document.tables:
        table = document.tables[0]
        if len(table.rows) < 2:
            return [], [{"row": 0, "error": "Word table has no data rows."}]

        headers = [_norm_header(cell.text) for cell in table.rows[0].cells]
        missing = [c for c in SECRETARY_REQUIRED_COLUMNS if c not in headers]
        if missing:
            return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

        for i, tr in enumerate(table.rows[1:], start=2):
            values = [cell.text.strip() for cell in tr.cells]
            row = {
                headers[idx]: (values[idx] if idx < len(values) else "")
                for idx in range(len(headers))
            }
            rows.append(row)
            if not row.get("name") or not row.get("id_number") or not row.get("county") or not row.get("date_of_birth"):
                errors.append({"row": i, "error": "Missing one or more required values: name, id_number, county, date_of_birth"})
        return rows, errors

    # Fallback format: comma-separated lines with header line.
    lines = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    if not lines:
        return [], [{"row": 0, "error": "Word document is empty."}]

    header_parts = [_norm_header(p) for p in lines[0].split(",")]
    missing = [c for c in SECRETARY_REQUIRED_COLUMNS if c not in header_parts]
    if missing:
        return [], [{"row": 0, "error": f"Missing required columns: {', '.join(missing)}"}]

    for i, line in enumerate(lines[1:], start=2):
        parts = [p.strip() for p in line.split(",")]
        row = {
            header_parts[idx]: (parts[idx] if idx < len(parts) else "")
            for idx in range(len(header_parts))
        }
        rows.append(row)
        if not row.get("name") or not row.get("id_number") or not row.get("county") or not row.get("date_of_birth"):
            errors.append({"row": i, "error": "Missing one or more required values: name, id_number, county, date_of_birth"})

    return rows, errors


class CountyBulkPlayerUploadView(APIView):
    """
    POST /api/v1/teams/county-players/bulk-upload/
    Organising Secretary bulk uploads players by name, id, county, DOB.
    Players are auto-allocated to county disciplines.
    """
    permission_classes = [IsCompetitionManager]

    @extend_schema(tags=["teams"], summary="Organising Secretary bulk upload county players via CSV/Excel/Word")
    def post(self, request):
        file = request.FILES.get("file")
        default_sport_type = str(request.data.get("sport_type", "")).strip()

        if not file:
            return Response({"detail": "No file uploaded."}, status=status.HTTP_400_BAD_REQUEST)

        filename = file.name.lower()
        if filename.endswith(".csv"):
            rows, parse_errors = _parse_secretary_csv(file)
        elif filename.endswith((".xlsx", ".xls")):
            rows, parse_errors = _parse_secretary_excel(file)
        elif filename.endswith(".docx"):
            rows, parse_errors = _parse_secretary_word(file)
        elif filename.endswith(".doc"):
            return Response({
                "detail": "Legacy .doc is not supported. Please save as .docx, .csv, or .xlsx and retry."
            }, status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response(
                {"detail": "Unsupported file format. Use .csv, .xlsx, or .docx"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if parse_errors:
            return Response({
                "detail": "File has validation errors. No players were created.",
                "errors": parse_errors,
                "total_rows": len(rows),
            }, status=status.HTTP_400_BAD_REQUEST)

        valid_sport_values = {choice[0] for choice in SportType.choices}
        if default_sport_type and default_sport_type not in valid_sport_values:
            return Response({
                "detail": "Invalid sport_type. Use a valid SportType value (e.g. handball_men, handball_women)."
            }, status=status.HTTP_400_BAD_REQUEST)

        county_lookup = {str(c.value).strip().lower(): str(c.value) for c in KenyaCounty}
        created = 0
        updated = 0
        row_errors = []

        for i, row in enumerate(rows, start=2):
            name_text = _extract_secretary_value(row, "name", "full_name", "player_name")
            national_id = _extract_secretary_value(row, "id_number", "id", "national_id", "national_id_number")
            county_raw = _extract_secretary_value(row, "county")
            dob_text = _extract_secretary_value(row, "date_of_birth", "dob", "date_of_birth_")
            row_sport_type = _extract_secretary_value(row, "discipline", "sport_type")

            sport_type = row_sport_type or default_sport_type
            if not sport_type:
                row_errors.append({
                    "row": i,
                    "error": "Discipline is required. Provide sport_type in request or discipline column in file.",
                })
                continue
            if sport_type not in valid_sport_values:
                row_errors.append({"row": i, "error": f"Invalid discipline/sport_type '{sport_type}'"})
                continue

            county_key = county_raw.strip().lower()
            county_value = county_lookup.get(county_key)
            if not county_value:
                row_errors.append({"row": i, "error": f"Invalid county '{county_raw}'"})
                continue

            parsed_dob = _parse_date(dob_text)
            if not parsed_dob:
                row_errors.append({"row": i, "error": f"Invalid date_of_birth '{dob_text}'"})
                continue

            first_name, last_name = _split_name(name_text)
            if not first_name:
                row_errors.append({"row": i, "error": "Invalid name value."})
                continue

            registration = CountyRegistration.objects.filter(county=county_value).first()
            if not registration:
                row_errors.append({
                    "row": i,
                    "error": f"County registration not found for '{county_value}'.",
                })
                continue

            discipline, _ = CountyDiscipline.objects.get_or_create(
                registration=registration,
                sport_type=sport_type,
            )

            try:
                player, was_created = CountyPlayer.objects.update_or_create(
                    national_id_number=national_id,
                    defaults={
                        "discipline": discipline,
                        "first_name": first_name,
                        "last_name": last_name,
                        "date_of_birth": parsed_dob,
                        # File format does not include phone; set placeholder for now.
                        "phone": "+254700000000",
                    },
                )
                if was_created:
                    created += 1
                else:
                    updated += 1
            except Exception as exc:
                row_errors.append({"row": i, "error": str(exc)})

        code = status.HTTP_201_CREATED if not row_errors else status.HTTP_207_MULTI_STATUS
        return Response({
            "detail": f"Bulk upload complete: {created} created, {updated} updated.",
            "created": created,
            "updated": updated,
            "total_rows": len(rows),
            "errors": row_errors,
        }, status=code)


class CountyBulkPlayerTemplateView(APIView):
    """
    GET /api/v1/teams/county-players/bulk-template/
    Download CSV template for Organising Secretary bulk county-player upload.
    """
    permission_classes = [IsCompetitionManager]

    @extend_schema(tags=["teams"], summary="Download Organising Secretary county-player bulk template")
    def get(self, request):
        response = HttpResponse(content_type="text/csv")
        response["Content-Disposition"] = 'attachment; filename="kyisa_county_player_upload_template.csv"'

        writer = csv.writer(response)
        writer.writerow(["name", "id_number", "county", "date_of_birth", "discipline"])
        writer.writerow(["John Doe", "12345678", "Makueni", "2008-03-15", "handball_men"])
        writer.writerow(["Jane Akinyi", "23456789", "Siaya", "2008-07-20", "handball_women"])
        return response
